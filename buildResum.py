import xlwt
import xlrd
import  os
from xlutils.copy import copy

import pandas as pd
import docx





# 创建一个workbook 设置编码
def addResum(base_massage,personal_summary,edu_Experience,item_Experience,other_honor,Student_ID):

    workbook = xlwt.Workbook(encoding='utf-8')
    # # 创建一个worksheet
    worksheet = workbook.add_sheet('My Worksheet')
    #
    # # 写入excel
    # # 参数对应 行, 列, 值
    worksheet.write(0, 0, label='基本信息')

    worksheet.write(0, 1, label=base_massage)

    worksheet.write(1, 0, label='个人总结')

    worksheet.write(1, 1, label=personal_summary)

    worksheet.write(2, 0, label='教育经历')

    worksheet.write(2, 1, label=edu_Experience)

    worksheet.write(3, 0, label='项目经历')

    worksheet.write(3, 1, label=item_Experience)

    worksheet.write(4, 0, label='证书/技能及其他')

    worksheet.write(4, 1, label=other_honor)

    workbook.save(r'static/Student_Resum/'+Student_ID+'/'+Student_ID+'.xls')

def updateResum(base_massage,personal_summary,edu_Experience,item_Experience,other_honor,Student_ID):
    readbook = xlrd.open_workbook(r'static/Student_Resum/'+Student_ID+'/'+Student_ID+'.xls')

    copybook = copy(readbook)

    writesheet = copybook.get_sheet(0)

    writesheet.write(0, 0, label='基本信息')

    writesheet.write(0, 1, label=base_massage)

    writesheet.write(1, 0, label='个人总结')

    writesheet.write(1, 1, label=personal_summary)

    writesheet.write(2, 0, label='教育经历')

    writesheet.write(2, 1, label=edu_Experience)

    writesheet.write(3, 0, label='项目经历')

    writesheet.write(3, 1, label=item_Experience)

    writesheet.write(4, 0, label='证书/技能及其他')

    writesheet.write(4, 1, label=other_honor)

    copybook.save(r'static/Student_Resum/'+Student_ID+'/'+Student_ID+'.xls')



def showResum(Student_ID):

    io = r'static/Student_Resum/'+Student_ID+'/'+Student_ID+'.xls'

    data = pd.read_excel(io, sheet_name ='My Worksheet',usecols = [0, 1],encoding='utf-8')

    pd.set_option('display.max_rows', 1000)

    pd.set_option('display.width', 1000)

    print(data)

    return data


def setnull(io):
    date = docx.Document()

    date.add_paragraph('评论内容')

    date.save(io)



def useResum(base_massage,personal_summary,edu_Experience,item_Experience,other_honor,Student_ID):

    io = r'static/Student_Resum/'+Student_ID+'/'+Student_ID+'.docx'

    student_resums ='基本信息：'+'\n'+base_massage+'\n'+'个人总结：'+'\n'+personal_summary+\
                    '\n'+'教育经历'+'\n'+edu_Experience+'\n'+'项目经历'+'\n'+item_Experience+'\n'+'证书及其他'+'\n'+other_honor

    date = docx.Document()

    date.add_paragraph(student_resums)

    date.save(io)

    return io


def makeDiscuss(io,discuss):

    # size = os.path.getsize(io)
    # if size>0:
    #     with open('io', 'w', encoding='utf-8')as f:
    #         f.write(discuss)
    # else:
    date = docx.Document(io)

    date.add_paragraph(discuss)

    date.save(io)

    io = str(io).split('/')[3]

    return io



def getResum(Student_ID):

    io = r'static/Student_Resum/' + Student_ID + '/' + Student_ID + '.xls'

    size = os.path.getsize(io)

    if size ==0:

        os.remove(io)

        workbook = xlwt.Workbook(encoding='utf-8')

        # # 创建一个worksheet
        worksheet = workbook.add_sheet('My Worksheet')

        worksheet.write(0, 0, label='无')

        workbook.save(r'static/Student_Resum/' + Student_ID + '/' + Student_ID + '.xls')

    else:
        workbook = xlrd.open_workbook(io)

        sheet =workbook.sheet_by_name('My Worksheet')

        nrows =sheet.nrows

        ncols =sheet.ncols

        base_massage= sheet.cell(0,1).value

        personal_summary=sheet.cell(1,1).value

        edu_Experience = sheet.cell(2, 1).value

        item_Experience = sheet.cell(3, 1).value

        other_honor = sheet.cell(4, 1).value

    return base_massage,personal_summary,edu_Experience,item_Experience,other_honor



def viewResum(filename):

    doc = docx.Document(filename)

    fullText = []

    for i in doc.paragraphs:  # 迭代docx文档里面的每一个段落

        fullText.append(i.text)  # 保存每一个段落的文本

    return '\n'.join(fullText)

